from flask import *
from openai import OpenAI
from docx import Document
from fpdf import FPDF
import os
import tempfile
from flask_cors import *

app = Flask(__name__)
CORS(app)
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def transcribe_audio(audio_file_path, language='en'):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file)
    return transcription.text

def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": "You are a highly skilled AI..."},
            {"role": "user", "content": transcription}
        ]
    )
    return response.choices[0].message.content

def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": "You are a proficient AI..."},
            {"role": "user", "content": transcription}
        ]
    )
    return response.choices[0].message.content

def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": "You are an AI expert..."},
            {"role": "user", "content": transcription}
        ]
    )
    return response.choices[0].message.content

def sentiment_analysis(transcription):
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": "As an AI with expertise..."},
            {"role": "user", "content": transcription}
        ]
    )
    return response.choices[0].message.content

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'transcription': transcription,
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()
    doc.save(filename)

def save_as_pdf(minutes, filename):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        pdf.set_font("Arial", 'B', size=14)
        pdf.cell(200, 10, txt=heading, ln=True, align='L')
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=value)
        pdf.ln(10)
    pdf.output(filename)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    filename = request.form['filename']
    format_choice = request.form['format_choice']
    temp_dir = tempfile.mkdtemp()
    audio_path = os.path.join(temp_dir, 'audio.mp3')
    file.save(audio_path)

    transcription = transcribe_audio(audio_path)
    minutes = meeting_minutes(transcription)

    if format_choice == 'docx':
        output_filename = os.path.join(temp_dir, filename + '.docx')
        save_as_docx(minutes, output_filename)
    elif format_choice == 'pdf':
        output_filename = os.path.join(temp_dir, filename + '.pdf')
        save_as_pdf(minutes, output_filename)
    else:
        return jsonify({'error': 'Invalid format choice'}), 400

    return send_file(output_filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
